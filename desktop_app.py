#!/usr/bin/env python3
"""
Word文档AI语法检查器 - PyQt6桌面版
现代化的原生桌面界面
"""

import sys
import json
import os
import time
import tempfile
from pathlib import Path
from typing import List, Dict
import pandas as pd

try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QLabel, QLineEdit, QPushButton, QComboBox, QTextEdit, QFileDialog,
        QProgressBar, QTabWidget, QFormLayout, QSpinBox, QDoubleSpinBox,
        QCheckBox, QListWidget, QListWidgetItem, QMessageBox, QSplitter,
        QGroupBox, QRadioButton, QButtonGroup, QScrollArea
    )
    from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
    from PyQt6.QtGui import QFont, QIcon, QPixmap
    from docx import Document
except ImportError as e:
    print(f"缺少必要的库: {e}")
    print("请运行: pip install PyQt6 python-docx litellm pandas openpyxl")
    sys.exit(1)

from utils.checker_core import process_paragraphs as core_process_paragraphs
from i18n import get_text


class ProcessingThread(QThread):
    """后台处理线程"""
    progress_updated = pyqtSignal(int, str)
    processing_finished = pyqtSignal(pd.DataFrame)
    error_occurred = pyqtSignal(str)

    def __init__(self, paragraphs, config):
        super().__init__()
        self.paragraphs = paragraphs
        self.config = config

    def run(self):
        """运行处理任务"""
        try:
            if not self.config.get("api_key"):
                lang = 'zh' if self.config.get('language', '中文') == '中文' else 'en'
                self.error_occurred.emit(get_text('API密钥不能为空', lang))
                return

            def callback(i: int, total: int, message: str):
                progress = int((i + 1) / total * 100)
                self.progress_updated.emit(progress, message)

            results = core_process_paragraphs(
                self.paragraphs, self.config, progress_callback=callback
            )
            df = pd.DataFrame(results)
            self.processing_finished.emit(df)
        except Exception as e:
            self.error_occurred.emit(str(e))


class MainWindow(QMainWindow):
    """主窗口"""

    def __init__(self):
        super().__init__()
        self.current_paragraphs = []
        self.current_file = None
        self.processing_thread = None
        self.current_language = 'zh'
        self.translatable_widgets = []
        self.init_ui()
        self.load_config()

    def tr(self, text: str) -> str:
        return get_text(text, self.current_language)

    def add_translatable(self, widget, text: str, attr: str = 'setText'):
        self.translatable_widgets.append((widget, text, attr))
        getattr(widget, attr)(self.tr(text))

    def apply_language(self, button=None, checked=False):
        self.current_language = 'zh' if self.chinese_radio.isChecked() else 'en'
        self.retranslate_ui()

    def retranslate_ui(self):
        self.setWindowTitle(self.tr('AI语法检查器'))
        for widget, text, attr in self.translatable_widgets:
            getattr(widget, attr)(self.tr(text))
        if self.current_file is None:
            self.file_label.setText(self.tr('未选择文件'))
        if not self.current_paragraphs:
            self.preview_text.setPlaceholderText(self.tr('文档预览将在这里显示...'))
        default_cn = get_text('语法检查结果.xlsx', 'zh')
        default_en = get_text('语法检查结果.xlsx', 'en')
        current_name = Path(self.output_path_input.text()).name
        if current_name in (default_cn, default_en):
            self.output_path_input.setText(str(Path.home() / self.tr('语法检查结果.xlsx')))
    
    def init_ui(self):
        """初始化界面"""
        self.setWindowTitle(self.tr("AI语法检查器"))
        self.setGeometry(100, 100, 1200, 800)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 创建主布局
        main_layout = QHBoxLayout(central_widget)
        
        # 创建分割器
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)
        
        # 左侧：配置面板
        left_widget = self.create_config_panel()
        splitter.addWidget(left_widget)
        
        # 右侧：主要操作面板
        right_widget = self.create_main_panel()
        splitter.addWidget(right_widget)
        
        # 设置分割器比例
        splitter.setSizes([400, 800])
    
    def create_config_panel(self) -> QWidget:
        """创建配置面板"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 标题
        title = QLabel()
        self.add_translatable(title, "⚙️ 配置设置")
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        layout.addWidget(title)
        
        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        
        # API设置组
        api_group = QGroupBox()
        self.add_translatable(api_group, "🔑 API设置", "setTitle")
        api_layout = QFormLayout(api_group)

        self.openai_key_input = QLineEdit()
        self.openai_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        openai_label = QLabel()
        self.add_translatable(openai_label, "OpenAI API Key:")
        api_layout.addRow(openai_label, self.openai_key_input)

        self.gemini_key_input = QLineEdit()
        self.gemini_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        gemini_label = QLabel()
        self.add_translatable(gemini_label, "Gemini API Key:")
        api_layout.addRow(gemini_label, self.gemini_key_input)
        
        scroll_layout.addWidget(api_group)
        
        # 模型设置组
        model_group = QGroupBox()
        self.add_translatable(model_group, "🤖 模型设置", "setTitle")
        model_layout = QFormLayout(model_group)
        
        self.provider_combo = QComboBox()
        self.provider_combo.addItems(["openai", "gemini"])
        self.provider_combo.currentTextChanged.connect(self.on_provider_changed)
        provider_label = QLabel()
        self.add_translatable(provider_label, "AI供应商:")
        model_layout.addRow(provider_label, self.provider_combo)

        self.model_combo = QComboBox()
        model_label = QLabel()
        self.add_translatable(model_label, "模型:")
        model_layout.addRow(model_label, self.model_combo)
        
        scroll_layout.addWidget(model_group)
        
        # 语言设置组
        language_group = QGroupBox()
        self.add_translatable(language_group, "🌐 语言设置", "setTitle")
        language_layout = QVBoxLayout(language_group)
        
        self.language_group = QButtonGroup()
        self.chinese_radio = QRadioButton()
        self.add_translatable(self.chinese_radio, "中文")
        self.english_radio = QRadioButton()
        self.add_translatable(self.english_radio, "English")
        self.chinese_radio.setChecked(True)

        self.language_group.addButton(self.chinese_radio, 0)
        self.language_group.addButton(self.english_radio, 1)
        self.language_group.buttonToggled.connect(self.apply_language)
        
        language_layout.addWidget(self.chinese_radio)
        language_layout.addWidget(self.english_radio)
        scroll_layout.addWidget(language_group)
        
        # 高级设置组
        advanced_group = QGroupBox()
        self.add_translatable(advanced_group, "🔧 高级设置", "setTitle")
        advanced_layout = QFormLayout(advanced_group)
        
        self.max_retries_spin = QSpinBox()
        self.max_retries_spin.setRange(1, 10)
        self.max_retries_spin.setValue(3)
        retries_label = QLabel()
        self.add_translatable(retries_label, "最大重试次数:")
        advanced_layout.addRow(retries_label, self.max_retries_spin)
        
        self.retry_delay_spin = QDoubleSpinBox()
        self.retry_delay_spin.setRange(0.1, 10.0)
        self.retry_delay_spin.setValue(1.0)
        self.retry_delay_spin.setSingleStep(0.1)
        retry_label = QLabel()
        self.add_translatable(retry_label, "重试延迟(秒):")
        advanced_layout.addRow(retry_label, self.retry_delay_spin)
        
        self.session_interval_spin = QSpinBox()
        self.session_interval_spin.setRange(1, 20)
        self.session_interval_spin.setValue(3)
        session_label = QLabel()
        self.add_translatable(session_label, "会话刷新间隔:")
        advanced_layout.addRow(session_label, self.session_interval_spin)
        
        scroll_layout.addWidget(advanced_group)
        
        # 配置文件操作
        config_group = QGroupBox()
        self.add_translatable(config_group, "💾 配置管理", "setTitle")
        config_layout = QVBoxLayout(config_group)

        save_config_btn = QPushButton()
        self.add_translatable(save_config_btn, "保存配置")
        save_config_btn.clicked.connect(self.save_config)
        config_layout.addWidget(save_config_btn)

        load_config_btn = QPushButton()
        self.add_translatable(load_config_btn, "加载配置")
        load_config_btn.clicked.connect(self.load_config_file)
        config_layout.addWidget(load_config_btn)
        
        scroll_layout.addWidget(config_group)
        
        scroll_area.setWidget(scroll_content)
        scroll_area.setWidgetResizable(True)
        layout.addWidget(scroll_area)
        
        return widget
    
    def create_main_panel(self) -> QWidget:
        """创建主操作面板"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 文件上传区域
        file_group = QGroupBox()
        self.add_translatable(file_group, "📁 文件操作", "setTitle")
        file_layout = QVBoxLayout(file_group)
        
        file_btn_layout = QHBoxLayout()
        self.select_file_btn = QPushButton()
        self.add_translatable(self.select_file_btn, "选择Word文档")
        self.select_file_btn.clicked.connect(self.select_word_file)
        file_btn_layout.addWidget(self.select_file_btn)

        self.file_label = QLabel(self.tr("未选择文件"))
        file_btn_layout.addWidget(self.file_label)
        file_btn_layout.addStretch()
        
        file_layout.addLayout(file_btn_layout)
        
        # 文档预览
        self.preview_text = QTextEdit()
        self.preview_text.setMaximumHeight(150)
        self.add_translatable(self.preview_text, "文档预览将在这里显示...", "setPlaceholderText")
        file_layout.addWidget(self.preview_text)
        
        layout.addWidget(file_group)
        
        # 输出设置
        output_group = QGroupBox()
        self.add_translatable(output_group, "📊 输出设置", "setTitle")
        output_layout = QFormLayout(output_group)
        
        output_btn_layout = QHBoxLayout()
        self.output_path_input = QLineEdit()
        self.output_path_input.setText(str(Path.home() / self.tr("语法检查结果.xlsx")))
        output_btn_layout.addWidget(self.output_path_input)
        
        browse_btn = QPushButton()
        self.add_translatable(browse_btn, "浏览")
        browse_btn.clicked.connect(self.browse_output_path)
        output_btn_layout.addWidget(browse_btn)

        excel_label = QLabel()
        self.add_translatable(excel_label, "Excel输出路径:")
        output_layout.addRow(excel_label, output_btn_layout)
        layout.addWidget(output_group)
        
        # 额外检查要求
        checks_group = QGroupBox()
        self.add_translatable(checks_group, "✅ 额外检查要求", "setTitle")
        checks_layout = QVBoxLayout(checks_group)
        
        add_check_layout = QHBoxLayout()
        self.new_check_input = QLineEdit()
        self.add_translatable(self.new_check_input, "输入新的检查要求...", "setPlaceholderText")
        add_check_layout.addWidget(self.new_check_input)

        add_check_btn = QPushButton()
        self.add_translatable(add_check_btn, "添加")
        add_check_btn.clicked.connect(self.add_check_requirement)
        add_check_layout.addWidget(add_check_btn)
        
        checks_layout.addLayout(add_check_layout)
        
        self.checks_list = QListWidget()
        checks_layout.addWidget(self.checks_list)
        
        remove_check_btn = QPushButton()
        self.add_translatable(remove_check_btn, "删除选中项")
        remove_check_btn.clicked.connect(self.remove_check_requirement)
        checks_layout.addWidget(remove_check_btn)
        
        layout.addWidget(checks_group)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel()
        self.status_label.setVisible(False)
        layout.addWidget(self.status_label)
        
        # 开始按钮
        self.start_btn = QPushButton()
        self.add_translatable(self.start_btn, "🚀 开始语法检查")
        self.start_btn.setMinimumHeight(50)
        self.start_btn.setFont(QFont("Arial", 12, QFont.Weight.Bold))
        self.start_btn.clicked.connect(self.start_processing)
        layout.addWidget(self.start_btn)

        credit_label = QLabel("Vibe coded with Claude and CodeX")
        credit_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(credit_label)

        return widget
    
    def on_provider_changed(self, provider):
        """供应商改变时更新模型列表"""
        self.model_combo.clear()
        
        if provider == "openai":
            models = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-4", "gpt-3.5-turbo"]
        elif provider == "gemini":
            models = ["gemini-pro", "gemini-pro-vision", "gemini-1.5-pro", "gemini-1.5-flash"]
        
        self.model_combo.addItems(models)
    
    def select_word_file(self):
        """选择Word文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("选择Word文档"), "", "Word文档 (*.docx)"
        )
        
        if file_path:
            self.file_label.setText(Path(file_path).name)
            self.load_word_document(file_path)
    
    def load_word_document(self, file_path):
        """加载Word文档"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            self.current_paragraphs = paragraphs
            
            # 显示预览
            preview_text = self.tr("文档包含 {} 个段落\n\n").format(len(paragraphs))
            for i, para in enumerate(paragraphs[:3]):
                preview_text += self.tr("段落 {}: {}...\n\n").format(i + 1, para[:100])

            self.preview_text.setText(preview_text)

        except Exception as e:
            QMessageBox.critical(self, self.tr("错误"), self.tr("读取Word文档失败: {}").format(e))
    
    def browse_output_path(self):
        """浏览输出路径"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("保存Excel文件"), "", "Excel文件 (*.xlsx)"
        )
        
        if file_path:
            self.output_path_input.setText(file_path)
    
    def add_check_requirement(self):
        """添加检查要求"""
        text = self.new_check_input.text().strip()
        if text:
            self.checks_list.addItem(text)
            self.new_check_input.clear()
    
    def remove_check_requirement(self):
        """删除检查要求"""
        current_row = self.checks_list.currentRow()
        if current_row >= 0:
            self.checks_list.takeItem(current_row)
    
    def get_additional_checks(self) -> List[str]:
        """获取额外检查要求列表"""
        checks = []
        for i in range(self.checks_list.count()):
            checks.append(self.checks_list.item(i).text())
        return checks
    
    def save_config(self):
        """保存配置"""
        config = {
            "provider": self.provider_combo.currentText(),
            "model": self.model_combo.currentText(),
            "language": "中文" if self.chinese_radio.isChecked() else "English",
            "max_retries": self.max_retries_spin.value(),
            "retry_delay": self.retry_delay_spin.value(),
            "session_refresh_interval": self.session_interval_spin.value(),
            "additional_checks": self.get_additional_checks()
        }
        
        try:
            with open("config.json", "w", encoding="utf-8") as f:
                json.dump(config, f, indent=2, ensure_ascii=False)
            QMessageBox.information(self, self.tr("成功"), self.tr("配置已保存到 config.json"))
        except Exception as e:
            QMessageBox.critical(self, self.tr("错误"), self.tr("保存配置失败: {}").format(e))
    
    def load_config(self):
        """加载配置"""
        try:
            if os.path.exists("config.json"):
                with open("config.json", "r", encoding="utf-8") as f:
                    config = json.load(f)
                
                self.openai_key_input.setText(os.getenv("OPENAI_API_KEY", config.get("openai_api_key", "")))
                self.gemini_key_input.setText(os.getenv("GEMINI_API_KEY", config.get("gemini_api_key", "")))
                
                provider = config.get("provider", "openai")
                self.provider_combo.setCurrentText(provider)
                self.on_provider_changed(provider)
                
                self.model_combo.setCurrentText(config.get("model", "gpt-3.5-turbo"))
                
                language = config.get("language", "中文")
                if language == "中文":
                    self.chinese_radio.setChecked(True)
                else:
                    self.english_radio.setChecked(True)
                self.apply_language()
                
                self.max_retries_spin.setValue(config.get("max_retries", 3))
                self.retry_delay_spin.setValue(config.get("retry_delay", 1.0))
                self.session_interval_spin.setValue(config.get("session_refresh_interval", 3))
                
                # 加载额外检查要求
                for check in config.get("additional_checks", []):
                    self.checks_list.addItem(check)
        
        except Exception as e:
            print(f"加载配置失败: {e}")
    
    def load_config_file(self):
        """从文件加载配置"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("选择配置文件"), "", "JSON文件 (*.json)"
        )
        
        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                # 这里可以复用load_config的逻辑
                QMessageBox.information(self, self.tr("成功"), self.tr("配置文件加载成功"))
            except Exception as e:
                QMessageBox.critical(self, self.tr("错误"), self.tr("加载配置文件失败: {}").format(e))
    
    def start_processing(self):
        """开始处理"""
        # 验证输入
        if not self.current_paragraphs:
            QMessageBox.warning(self, self.tr("警告"), self.tr("请先选择Word文档"))
            return

        provider = self.provider_combo.currentText()
        api_key = (os.getenv("OPENAI_API_KEY") if provider == "openai"
                  else os.getenv("GEMINI_API_KEY"))
        if not api_key:
            api_key = (self.openai_key_input.text() if provider == "openai"
                      else self.gemini_key_input.text())
        if not api_key:
            QMessageBox.warning(self, self.tr("警告"), self.tr("请输入{} API密钥").format(provider.upper()))
            return

        # 准备配置
        config = {
            "provider": provider,
            "model": self.model_combo.currentText(),
            "api_key": api_key,
            "language": "中文" if self.chinese_radio.isChecked() else "English",
            "session_refresh_interval": self.session_interval_spin.value(),
            "additional_checks": self.get_additional_checks()
        }
        
        # 禁用开始按钮，显示进度条
        self.start_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.status_label.setVisible(True)
        
        # 创建并启动处理线程
        self.processing_thread = ProcessingThread(self.current_paragraphs, config)
        self.processing_thread.progress_updated.connect(self.update_progress)
        self.processing_thread.processing_finished.connect(self.on_processing_finished)
        self.processing_thread.error_occurred.connect(self.on_processing_error)
        self.processing_thread.start()
    
    def update_progress(self, value, message):
        """更新进度"""
        self.progress_bar.setValue(value)
        self.status_label.setText(message)
    
    def on_processing_finished(self, result_df):
        """处理完成"""
        try:
            # 保存Excel文件
            output_path = self.output_path_input.text()
            result_df.to_excel(output_path, index=False, engine='openpyxl')
            
            QMessageBox.information(
                self,
                self.tr("完成"),
                self.tr("语法检查完成！\n结果已保存到: {}\n共处理 {} 个段落").format(
                    output_path, len(result_df)
                ),
            )
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("错误"), self.tr("保存Excel文件失败: {}").format(e))
        
        finally:
            # 重置界面
            self.start_btn.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.status_label.setVisible(False)
    
    def on_processing_error(self, error_message):
        """处理错误"""
        QMessageBox.critical(self, self.tr("错误"), self.tr("处理过程中出现错误: {}").format(error_message))
        
        # 重置界面
        self.start_btn.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.status_label.setVisible(False)


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')  # 使用现代化样式
    
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
